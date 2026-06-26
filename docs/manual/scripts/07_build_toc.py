#!/usr/bin/env python3
"""
VITMATERNA — Índice (TOC) con números de página REALES, determinista.

python-docx no calcula los números de página del campo TOC (los hace Word al
abrir). Para un PDF/DOCX entregable con índice completo:

  1. Se construye el DOCX con un MARCADOR de índice (párrafo "[[TOC]]").
  2. Se exporta a PDF (LibreOffice) — pasada 1.
  3. Se busca, con pdftotext por página, en qué página cae cada título.
  4. Se REEMPLAZA el marcador por una tabla de índice (título … página) y se
     reexporta el PDF — pasada 2.

Uso: python3 07_build_toc.py <ruta_docx> <lista_de_titulos.json>
"""
import sys
import os
import json
import re
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x0C, 0x81, 0x74)
INK = RGBColor(0x16, 0x24, 0x2B)


def slug(title):
    return "bm_" + re.sub(r"[^a-zA-Z0-9]", "_", title)[:38]


def add_bookmark_to_paragraph(p, name, bmid):
    """Envuelve el contenido del párrafo con un bookmark (destino de enlace)."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bmid)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bmid))
    p._p.insert(0, start)
    p._p.append(end)


def add_internal_hyperlink(paragraph, anchor, runs_specs):
    """Añade un hipervínculo interno (a un bookmark) con varios runs."""
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    for text, size, bold, color in runs_specs:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), "Carlito"); rFonts.set(qn("w:hAnsi"), "Carlito"); rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
        if bold:
            rPr.append(OxmlElement("w:b"))
        col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
        r.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
        hl.append(r)
    paragraph._p.append(hl)


def docx_to_pdf(docx, outdir):
    prof = subprocess.run(["mktemp", "-d"], capture_output=True, text=True).stdout.strip()
    subprocess.run(["soffice", "--headless", "--norestore",
                    f"-env:UserInstallation=file://{prof}",
                    "--convert-to", "pdf", "--outdir", outdir, docx],
                   capture_output=True, timeout=120)
    return docx.replace(".docx", ".pdf")


def page_of_titles(pdf, titles):
    """Devuelve {titulo: pagina} usando pdftotext por página."""
    txt = subprocess.run(["pdftotext", "-layout", pdf, "-"], capture_output=True, text=True).stdout
    pages = txt.split("\f")
    result = {}
    for t in titles:
        for i, pg in enumerate(pages, start=1):
            if t in pg and t not in result:
                result[t] = i
                break
    return result


def add_dotted_toc_row(doc_table, title, page, level, accent_hex):
    row = doc_table.add_row()
    c0, c1 = row.cells
    p = c0.paragraphs[0]
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.2
    indent = "        " if level >= 2 else ""
    size = 11 if level == 1 else 10
    color = accent_hex if level == 1 else "16242B"
    # Entrada como HIPERVÍNCULO interno al bookmark del título.
    add_internal_hyperlink(p, slug(title), [(f"{indent}{title}", size, level == 1, color)])
    p2 = c1.paragraphs[0]; p2.alignment = 2  # derecha
    # El número también enlaza al mismo destino.
    add_internal_hyperlink(p2, slug(title), [(str(page), 10, False, "16242B")])


def inject_heading_bookmarks(doc, titles):
    """Inserta un bookmark en el primer párrafo cuyo texto coincide con cada título."""
    want = {t: slug(t) for t in titles}
    bmid = 1000
    seen = set()
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt in want and txt not in seen and p.style.name.startswith("Heading"):
            add_bookmark_to_paragraph(p, want[txt], bmid)
            bmid += 1; seen.add(txt)


def replace_marker_with_toc(docx, toc_map, titles_levels, accent_hex):
    doc = Document(docx)
    inject_heading_bookmarks(doc, [t for t, _ in titles_levels])
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "[[TOC]]":
            for r in list(p.runs):
                r.text = ""
            tbl = doc.add_table(rows=0, cols=2)
            tbl.allow_autofit = True
            p._p.addnext(tbl._tbl)
            for title, level in titles_levels:
                if title in toc_map:
                    add_dotted_toc_row(tbl, title, toc_map[title], level, accent_hex)
            break
    doc.save(docx)


def main():
    docx = sys.argv[1]
    titles_levels = json.loads(open(sys.argv[2], encoding="utf-8").read())
    accent_hex = sys.argv[3] if len(sys.argv) > 3 else "0C8174"
    outdir = os.path.dirname(docx)
    titles = [t for t, _ in titles_levels]
    # pasada 1
    pdf = docx_to_pdf(docx, outdir)
    toc_map = page_of_titles(pdf, titles)
    # reemplazar marcador (con bookmarks + hipervínculos clickeables)
    replace_marker_with_toc(docx, toc_map, titles_levels, accent_hex)
    # pasada 2 (con el índice ya numerado y enlazado)
    docx_to_pdf(docx, outdir)
    print(f"✅ TOC con {len(toc_map)} entradas numeradas en {os.path.basename(docx)}")
    for t, _ in titles_levels:
        if t in toc_map:
            print(f"   {toc_map[t]:>3}  {t}")


if __name__ == "__main__":
    main()
