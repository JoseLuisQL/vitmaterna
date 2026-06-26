#!/usr/bin/env python3
"""
VITMATERNA — Generador de manual por ROL (obstetra | admin) en DOCX.

Lee manifest/<rol>.content.json (textos + qué captura va en cada sección) y las
capturas anotadas en assets/<rol>_annotated/, y arma un .docx profesional con la
misma estética que el de la gestante: portada, créditos, índice, intro,
requisitos, convenciones, acceso, secciones por tarea (pasos + captura), 
troubleshooting, FAQ, glosario y contacto.

Uso: python3 03b_build_role.py <obstetra|admin>
"""
import os
import sys
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROLE = sys.argv[1] if len(sys.argv) > 1 else "obstetra"
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPO = os.path.dirname(os.path.dirname(BASE))
IMG = os.path.join(BASE, "assets", f"{ROLE}_annotated")
LOGO = os.path.join(REPO, "vitmaterna_logo.png")
CONTENT = json.load(open(os.path.join(BASE, "manifest", f"{ROLE}.content.json"), encoding="utf-8"))
OUT = os.path.join(BASE, "build", f"manual_usuario_{ROLE}_vitmaterna_movil.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

ACCENT = RGBColor.from_string(CONTENT["accent"])
INK = RGBColor(0x16, 0x24, 0x2B)
GRAY = RGBColor(0x56, 0x68, 0x73)
RED = RGBColor(0xD6, 0x45, 0x45)
AMBER = RGBColor(0xB0, 0x7A, 0x14)
FIG = {"n": 0}
AUTHOR = "CRISTHIAN RODRIGO BERROCAL SALAZAR"
PREFIX = {"obstetra": "B", "admin": "C"}.get(ROLE, "X")


def set_cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def _field(run, instr):
    def el(tag, **attrs):
        e = OxmlElement(tag)
        for k, v in attrs.items():
            e.set(qn(k), v)
        return e
    run._r.append(el("w:fldChar", **{"w:fldCharType": "begin"}))
    it = el("w:instrText", **{"xml:space": "preserve"}); it.text = instr; run._r.append(it)
    run._r.append(el("w:fldChar", **{"w:fldCharType": "separate"}))
    run._r.append(el("w:fldChar", **{"w:fldCharType": "end"}))


def add_toc(doc):
    doc.add_paragraph("[[TOC]]")


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Carlito"; st.font.size = Pt(11); st.font.color.rgb = INK
    pf = st.paragraph_format
    pf.line_spacing = 1.4; pf.space_after = Pt(8); pf.space_before = Pt(0)
    specs = {
        "Heading 1": (16, ACCENT, 18, 8),
        "Heading 2": (13, INK, 14, 4),
        "Heading 3": (11.5, ACCENT, 10, 2),
    }
    for name, (size, color, before, after) in specs.items():
        s = doc.styles[name]
        s.font.name = "Carlito"; s.font.size = Pt(size); s.font.color.rgb = color; s.font.bold = True; s.font.italic = False
        pfh = s.paragraph_format
        pfh.space_before = Pt(before); pfh.space_after = Pt(after); pfh.line_spacing = 1.15; pfh.keep_with_next = True
    try:
        ln = doc.styles["List Number"]
        ln.font.size = Pt(11); ln.paragraph_format.space_after = Pt(6); ln.paragraph_format.line_spacing = 1.3
    except KeyError:
        pass


def caption(doc, title):
    FIG["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Figura {PREFIX}-{FIG['n']}. {title}"); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY


def shot(doc, fid, cap):
    path = os.path.join(IMG, f"{fid}.png")
    if not os.path.exists(path):
        doc.add_paragraph(f"[Captura {fid} no encontrada]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True; p.paragraph_format.keep_together = True; p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(path, height=Cm(11.3))
    caption(doc, cap)


def _cell_pad(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{side}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)


def note(doc, text, kind="nota"):
    colors = {"nota": ("EEF2F3", ACCENT, "Nota"), "aviso": ("FBF4E5", AMBER, "Importante"),
              "alerta": ("FBEDED", RED, "Atención")}
    fill, fg, label = colors[kind]
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.autofit = False
    c = tbl.cell(0, 0); set_cell_bg(c, fill); _cell_pad(c)
    try:
        c.width = Cm(16.6)
    except Exception:
        pass
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.3
    r = p.add_run(f"{label}.  "); r.bold = True; r.font.color.rgb = fg; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.font.size = Pt(10); r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def steps(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.3
        p.add_run(it)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.3
        p.add_run(it)


def _accent_rule(p, color):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), CONTENT["accent"])
    pbdr.append(bottom); pPr.append(pbdr)


def h1(doc, t):
    p = doc.add_heading(t, level=1); _accent_rule(p, CONTENT["accent"]); return p
def h2(doc, t): return doc.add_heading(t, level=2)
def h3(doc, t): return doc.add_heading(t, level=3)
def para(doc, t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.add_run(t); return p
def pb(doc): doc.add_page_break()
def section_gap(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)


def cover(doc):
    if os.path.exists(LOGO):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO, width=Cm(5))
    doc.add_paragraph()
    for txt, sz, col, bold in [
        ("VITMATERNA", 40, ACCENT, True),
        ("Manual de Usuario · Aplicación Móvil", 20, INK, False),
        (CONTENT["title"], 16, GRAY, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = col; r.bold = bold
    for _ in range(5):
        doc.add_paragraph()
    for txt, sz, col in [
        ("Plataforma de Salud Prenatal", 13, INK),
        ("Centro de Salud Talavera — Andahuaylas, Apurímac", 12, GRAY),
        ("Versión 1.0 · 2026", 11, GRAY),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = col
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Autor: "); r.font.size = Pt(11); r.font.color.rgb = GRAY
    r2 = p.add_run(AUTHOR); r2.font.size = Pt(11); r2.bold = True; r2.font.color.rgb = INK
    pb(doc)


def footer(doc):
    p = doc.sections[0].footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"VITMATERNA · {CONTENT['title']}     ·     Página "); r.font.size = Pt(8); r.font.color.rgb = GRAY
    rp = p.add_run(); rp.font.size = Pt(8); rp.font.color.rgb = GRAY; _field(rp, "PAGE")
    rmid = p.add_run(" de "); rmid.font.size = Pt(8); rmid.font.color.rgb = GRAY
    rt = p.add_run(); rt.font.size = Pt(8); rt.font.color.rgb = GRAY; _field(rt, "NUMPAGES")


def header(doc):
    p = doc.sections[0].header.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("VITMATERNA"); r.font.size = Pt(9); r.bold = True; r.font.color.rgb = ACCENT
    r2 = p.add_run("  ·  Salud Prenatal"); r2.font.size = Pt(9); r2.font.color.rgb = GRAY


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2); s.top_margin = s.bottom_margin = Cm(2.0)
    style_base(doc); footer(doc); header(doc); cover(doc)

    h1(doc, "Créditos y confidencialidad")
    pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.add_run("Autor del manual: ").bold = True; pa.add_run(AUTHOR)
    para(doc, "Este documento es el manual de usuario oficial de la aplicación móvil VITMATERNA, plataforma "
              "de seguimiento de salud prenatal. Se entrega con fines de capacitación y uso del sistema.")
    para(doc, "La información que aparece en las capturas corresponde a datos de demostración y no representa a "
              "personas reales. El manejo de datos de salud se realiza conforme a la normativa de protección de "
              "datos personales vigente.")
    note(doc, "Este documento está diseñado como guía de consulta; se recomienda utilizar el índice para acceder directamente a la sección deseada.")
    pb(doc)

    h1(doc, "Índice de contenidos"); add_toc(doc); pb(doc)

    h1(doc, "1. Introducción")
    h2(doc, "1.1 ¿Qué es VITMATERNA?"); para(doc, CONTENT["intro_que"])
    h2(doc, "1.2 Propósito de este manual"); para(doc, CONTENT["intro_quien"])
    h2(doc, "1.3 Contenido"); para(doc, CONTENT["intro_alcance"])

    h1(doc, "2. Antes de empezar")
    h2(doc, "2.1 Requisitos")
    para(doc, "Para utilizar la aplicación es necesario contar con:")
    bullets(doc, [
        "Un teléfono móvil con sistema operativo Android o iOS (iPhone).",
        "Conexión a internet (datos móviles o red wifi).",
        "Las credenciales de acceso: número de DNI y contraseña.",
    ])
    h2(doc, "2.2 Convenciones de este manual")
    para(doc, "En las imágenes verás marcas numeradas y recuadros que indican exactamente de qué elemento se "
              "habla y en qué orden usarlo:")
    tbl = doc.add_table(rows=3, cols=2); tbl.style = "Light Grid Accent 1"
    for i, (a, b) in enumerate([
        ("① ② ③", "Marca de paso: el número indica el orden en que debes tocar cada elemento."),
        ("Recuadro rojo", "Resalta el botón, la tarjeta o el campo que se está explicando."),
        ("Nota / Importante / Atención", "Avisos con información útil o advertencias a tener en cuenta."),
    ]):
        tbl.cell(i, 0).paragraphs[0].add_run(a).bold = True
        tbl.cell(i, 1).paragraphs[0].add_run(b)
    section_gap(doc)

    h1(doc, "3. Acceso al sistema")
    para(doc, "Al abrir la aplicación se muestra la pantalla de inicio de sesión. Ingresa tus credenciales "
              "para acceder a tu cuenta:")
    steps(doc, [
        "En el campo “DNI”, escribe tu número de documento.",
        "En el campo “Contraseña”, escribe tu contraseña.",
        "Toca el botón “Iniciar Sesión” para ingresar.",
    ])
    shot(doc, "A0", "Pantalla de inicio de sesión.")
    note(doc, "Durante el primer ingreso se mostrará el recorrido guiado “Conoce tu app”, que puede repetirse "
              "en cualquier momento desde el perfil.")
    section_gap(doc)

    for sec in CONTENT["sections"]:
        # Nivel de título según la profundidad del número: "4" -> H1, "4.1" -> H2
        num = sec["h"].split(" ")[0]
        if num.count(".") >= 1:
            h2(doc, sec["h"])
        else:
            h1(doc, sec["h"])
        if sec.get("p"): para(doc, sec["p"])
        if sec.get("steps"): steps(doc, sec["steps"])
        if sec.get("shot"): shot(doc, sec["shot"], sec.get("cap", ""))
        if sec.get("note"): note(doc, sec["note"], sec.get("note_kind", "nota"))

    section_gap(doc)
    h1(doc, "Solución de problemas")
    for q, a in CONTENT["troubleshooting"]:
        h3(doc, q); para(doc, a)

    h1(doc, "Preguntas frecuentes")
    for q, a in CONTENT["faq"]:
        h3(doc, q); para(doc, a)

    h1(doc, "Glosario")
    for t, d in CONTENT["glossary"]:
        p = doc.add_paragraph(); p.add_run(f"{t}: ").bold = True; p.add_run(d)

    h1(doc, "Soporte y contacto")
    para(doc, "Para soporte de la aplicación, comunícate con el área de sistemas de tu establecimiento. Para "
              "emergencias clínicas, acude al establecimiento de salud.")
    p = doc.add_paragraph(); p.add_run("Centro de Salud Talavera\n").bold = True
    p.add_run("Andahuaylas, Apurímac\nTeléfono: 083 – 421800")

    cp = doc.core_properties
    cp.author = AUTHOR
    cp.title = f"Manual de Usuario VITMATERNA — {CONTENT['title']}"
    cp.category = "Manual de Usuario"
    doc.save(OUT)
    print(f"✅ DOCX {ROLE}: {OUT}  ·  figuras: {FIG['n']}  ·  autor: {AUTHOR}")


if __name__ == "__main__":
    build()
