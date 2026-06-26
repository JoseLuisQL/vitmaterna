#!/usr/bin/env python3
"""
VITMATERNA — Generador del Manual de Usuario (Gestante) en DOCX.

Cobertura completa: pantallas principales + modales + formularios, con pasos
numerados que se corresponden con las marcas ①②③ de las capturas (medidas
pixel-perfect). Estructura profesional: portada, créditos, índice, intro,
requisitos, convenciones, acceso, secciones por tarea, troubleshooting, FAQ,
glosario, contacto.

Reproducible: `python3 03_build_docx.py`.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPO = os.path.dirname(os.path.dirname(BASE))
IMG = os.path.join(BASE, "assets", "screens_annotated")
LOGO = os.path.join(REPO, "vitmaterna_logo.png")
OUT = os.path.join(BASE, "build", "manual_usuario_gestante_vitmaterna_movil.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

TEAL = RGBColor(0x0C, 0x81, 0x74)
INK = RGBColor(0x16, 0x24, 0x2B)
GRAY = RGBColor(0x56, 0x68, 0x73)
RED = RGBColor(0xD6, 0x45, 0x45)
AMBER = RGBColor(0xB0, 0x7A, 0x14)
FIG = {"n": 0}
AUTHOR = "CRISTHIAN RODRIGO BERROCAL SALAZAR"


def set_cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def _field(run, instr):
    """Inserta un campo de Word COMPLETO (begin/instr/separate/end) en `run`.
    A diferencia de fldSimple, LibreOffice sí lo calcula al exportar (números
    de página reales)."""
    def el(tag, **attrs):
        e = OxmlElement(tag)
        for k, v in attrs.items():
            e.set(qn(k), v)
        return e
    begin = el("w:fldChar", **{"w:fldCharType": "begin"})
    it = el("w:instrText", **{"xml:space": "preserve"}); it.text = instr
    sep = el("w:fldChar", **{"w:fldCharType": "separate"})
    end = el("w:fldChar", **{"w:fldCharType": "end"})
    run._r.append(begin); run._r.append(it); run._r.append(sep); run._r.append(end)


def add_toc(doc):
    # Marcador que 07_build_toc.py reemplaza por una tabla de índice con números
    # de página reales (dos pasadas de render).
    doc.add_paragraph("[[TOC]]")


def style_base(doc):
    # Cuerpo: interlineado 1.4, espacio entre párrafos, lectura cómoda.
    st = doc.styles["Normal"]
    st.font.name = "Carlito"; st.font.size = Pt(11); st.font.color.rgb = INK
    pf = st.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)
    # Títulos: jerarquía clara con espacios de respiración y separación visual.
    specs = {
        "Heading 1": (16, TEAL, 18, 8, True),
        "Heading 2": (13, INK, 14, 4, True),
        "Heading 3": (11.5, TEAL, 10, 2, True),
    }
    for name, (size, color, before, after, keep) in specs.items():
        s = doc.styles[name]
        s.font.name = "Carlito"; s.font.size = Pt(size); s.font.color.rgb = color; s.font.bold = True
        s.font.italic = False
        pfh = s.paragraph_format
        pfh.space_before = Pt(before); pfh.space_after = Pt(after)
        pfh.line_spacing = 1.15; pfh.keep_with_next = True
    # Lista numerada: un poco más de aire entre pasos.
    try:
        ln = doc.styles["List Number"]
        ln.font.size = Pt(11); ln.paragraph_format.space_after = Pt(6); ln.paragraph_format.line_spacing = 1.3
    except KeyError:
        pass


def caption(doc, title):
    FIG["n"] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Figura A-{FIG['n']}. {title}"); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY


def shot(doc, fid, cap, height=11.3):
    path = os.path.join(IMG, f"{fid}.png")
    if not os.path.exists(path):
        doc.add_paragraph(f"[Captura {fid} no encontrada]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True; p.paragraph_format.keep_together = True; p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(path, height=Cm(height))
    caption(doc, cap)


def _cell_pad(cell, top=80, bottom=80, left=120, right=120):
    """Padding interno de una celda (en twips) para que el cuadro respire."""
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{side}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)


def note(doc, text, kind="nota"):
    colors = {"nota": ("E7F4F2", TEAL, "Nota"), "aviso": ("FBF4E5", AMBER, "Importante"),
              "alerta": ("FBEDED", RED, "Atención")}
    fill, fg, label = colors[kind]
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    c = tbl.cell(0, 0); set_cell_bg(c, fill); _cell_pad(c)
    try:
        c.width = Cm(16.6)
    except Exception:
        pass
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.3
    r = p.add_run(f"{label}.  "); r.bold = True; r.font.color.rgb = fg; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.font.size = Pt(10); r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)  # respiro tras el cuadro


def steps(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.3
        p.add_run(it)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.3
        p.add_run(it)


def _accent_rule(p, color=None):
    """Línea fina de color bajo un título (borde inferior del párrafo)."""
    color = color or "0C8174"
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)


def h1(doc, t):
    p = doc.add_heading(t, level=1)
    _accent_rule(p, "0C8174")
    return p
def h2(doc, t): return doc.add_heading(t, level=2)
def h3(doc, t): return doc.add_heading(t, level=3)


def para(doc, t, size=11):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(t).font.size = Pt(size); return p


def pb(doc): doc.add_page_break()
def section_gap(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)


def cover(doc):
    if os.path.exists(LOGO):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO, width=Cm(5))
    doc.add_paragraph()
    for txt, sz, col, bold in [
        ("VITMATERNA", 40, TEAL, True),
        ("Manual de Usuario · Aplicación Móvil", 20, INK, False),
        ("Guía para la Gestante", 16, GRAY, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = col; r.bold = bold
    for _ in range(5):
        doc.add_paragraph()
    for txt, sz, col in [
        ("Plataforma de Salud Prenatal", 13, INK),
        ("Centro de Salud Talavera — Andahuaylas, Apurímac", 12, GRAY),
        ("Versión 1.0 · 2026", 11, GRAY),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = col
    # Autor
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Autor: "); r.font.size = Pt(11); r.font.color.rgb = GRAY
    r2 = p.add_run(AUTHOR); r2.font.size = Pt(11); r2.bold = True; r2.font.color.rgb = INK
    pb(doc)


def footer(doc):
    p = doc.sections[0].footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VITMATERNA · Manual de Usuario (Gestante)     ·     Página "); r.font.size = Pt(8); r.font.color.rgb = GRAY
    rp = p.add_run(); rp.font.size = Pt(8); rp.font.color.rgb = GRAY; _field(rp, "PAGE")
    rmid = p.add_run(" de "); rmid.font.size = Pt(8); rmid.font.color.rgb = GRAY
    rt = p.add_run(); rt.font.size = Pt(8); rt.font.color.rgb = GRAY; _field(rt, "NUMPAGES")


def header(doc):
    # Encabezado discreto con el nombre del producto + acento de marca.
    p = doc.sections[0].header.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("VITMATERNA"); r.font.size = Pt(9); r.bold = True; r.font.color.rgb = TEAL
    r2 = p.add_run("  ·  Salud Prenatal"); r2.font.size = Pt(9); r2.font.color.rgb = GRAY


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2); s.top_margin = s.bottom_margin = Cm(2.0)
    style_base(doc); footer(doc); header(doc); cover(doc)

    # Créditos
    h1(doc, "Créditos y confidencialidad")
    pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.add_run("Autor del manual: ").bold = True; pa.add_run(AUTHOR)
    para(doc, "Este documento es el manual de usuario oficial de la aplicación móvil VITMATERNA, plataforma "
              "de seguimiento de salud prenatal. Se entrega con fines de capacitación y uso del sistema.")
    para(doc, "La información de pacientes que aparece en las capturas corresponde a datos de demostración y "
              "no representa a personas reales. El manejo de datos de salud se realiza conforme a la normativa "
              "de protección de datos personales vigente.")
    note(doc, "Este documento está diseñado como guía de consulta; se recomienda utilizar el índice para acceder directamente a la sección deseada.")
    pb(doc)

    # Índice
    h1(doc, "Índice de contenidos"); add_toc(doc); pb(doc)

    # 1. Introducción
    h1(doc, "1. Introducción")
    h2(doc, "1.1 ¿Qué es VITMATERNA?")
    para(doc, "VITMATERNA es una aplicación móvil de acompañamiento prenatal. Permite a la gestante seguir la "
              "evolución de su embarazo semana a semana, consultar y confirmar sus controles prenatales, "
              "registrar la toma de sus suplementos, comunicarse con su obstetra, acceder a contenido educativo "
              "y reportar signos de alarma o solicitar ayuda en caso de emergencia.")
    h2(doc, "1.2 Propósito de este manual")
    para(doc, "Este manual tiene como objetivo guiar a la gestante en el uso correcto de la aplicación. Describe "
              "cada función mediante instrucciones paso a paso acompañadas de capturas de pantalla, de modo que "
              "la usuaria pueda completar sus tareas con autonomía y seguridad.")
    h2(doc, "1.3 Contenido")
    para(doc, "El documento abarca el acceso al sistema, la pantalla de inicio, la gestión de citas, el "
              "seguimiento del tratamiento, la comunicación por chat, el contenido educativo, el reporte de "
              "signos de alarma y la administración del perfil personal.")

    # 2. Antes de empezar
    h1(doc, "2. Antes de empezar")
    h2(doc, "2.1 Requisitos")
    para(doc, "Para utilizar la aplicación es necesario contar con:")
    bullets(doc, [
        "Un teléfono móvil con sistema operativo Android o iOS (iPhone).",
        "Conexión a internet (datos móviles o red wifi).",
        "Las credenciales de acceso: número de DNI y contraseña.",
    ])
    h2(doc, "2.2 Instalación")
    para(doc, "Descarga la aplicación “VITMATERNA” desde la tienda de tu dispositivo (Play Store en Android o "
              "App Store en iPhone) e instálala. Una vez instalada, abre la aplicación tocando su ícono.")
    h2(doc, "2.3 Convenciones de este manual")
    para(doc, "En las imágenes verás marcas numeradas y recuadros que indican exactamente de qué elemento se "
              "habla y en qué orden usarlo:")
    tbl = doc.add_table(rows=3, cols=2); tbl.style = "Light Grid Accent 1"
    for i, (a, b) in enumerate([
        ("① ② ③", "Marca de paso: el número indica el orden en que debes tocar cada elemento de la pantalla."),
        ("Recuadro rojo", "Resalta el botón, la tarjeta o el campo que se está explicando."),
        ("Nota / Importante / Atención", "Avisos con información útil o advertencias a tener en cuenta."),
    ]):
        tbl.cell(i, 0).paragraphs[0].add_run(a).bold = True
        tbl.cell(i, 1).paragraphs[0].add_run(b)
    section_gap(doc)

    # 3. Acceso
    h1(doc, "3. Acceso al sistema")
    h2(doc, "3.1 Iniciar sesión")
    para(doc, "Al abrir la aplicación se muestra la pantalla de inicio de sesión. Ingresa tus credenciales "
              "para acceder a tu cuenta:")
    steps(doc, [
        "En el campo “DNI”, escribe tu número de documento.",
        "En el campo “Contraseña”, escribe tu contraseña.",
        "Toca el botón “Iniciar Sesión” para ingresar.",
    ])
    shot(doc, "A0", "Pantalla de inicio de sesión.")
    note(doc, "Si olvidaste tu contraseña, toca “¿Olvidaste tu contraseña?” o solicita ayuda a tu obstetra.")
    h2(doc, "3.2 Recorrido guiado “Conoce tu app”")
    para(doc, "Durante el primer ingreso, la aplicación ofrece un recorrido guiado que presenta sus principales "
              "funciones. Puede repetirse en cualquier momento desde “Mi perfil” → “Conoce tu app”.")
    section_gap(doc)

    # 4. Inicio
    h1(doc, "4. Pantalla de inicio")
    para(doc, "Es la primera pantalla al entrar. Te muestra un resumen de tu embarazo del día.")
    steps(doc, [
        "① “Tu embarazo”: la cinta muestra tu semana de gestación, el trimestre y tu nivel de riesgo.",
        "② “Próxima Cita”: el día, la hora y el estado de tu siguiente control.",
        "③ “Tratamiento del Día”: cuántas vitaminas o pastillas debes tomar hoy y tu avance.",
        "④ Botón de menú (☰): abre el menú lateral para ir a Educación, Perfil y más.",
    ])
    shot(doc, "A1", "Pantalla de inicio de la gestante.")
    note(doc, "La semana de embarazo y el avance se calculan solos a partir de tu fecha de última regla (FUM).")

    h2(doc, "4.1 Confirmar la asistencia a tu cita")
    steps(doc, ["① En la tarjeta “Próxima Cita”, toca “Confirmar asistencia” para avisar que asistirás. "
                "Tu obstetra recibirá la confirmación."])
    shot(doc, "A1c", "Botón para confirmar la asistencia a la próxima cita.")

    h2(doc, "4.2 Acciones rápidas")
    para(doc, "En la parte inferior del inicio tienes tres accesos directos:")
    steps(doc, [
        "① “Reportar”: informa un signo de alarma (síntoma) a tu obstetra.",
        "② “Emergencia”: envía una alerta de auxilio con tu ubicación.",
        "③ “Educación”: abre la biblioteca de contenido para tu embarazo.",
    ])
    shot(doc, "A1b", "Acciones rápidas del inicio.")
    note(doc, "Al tocar “Emergencia” se abre una ventana de confirmación con el botón “Enviar ahora”. Solo se "
              "envía la alerta cuando lo confirmas; también puedes “Cancelar”.", "alerta")

    h2(doc, "4.3 El menú lateral")
    para(doc, "Toca el botón de menú (☰) del inicio para abrir el menú lateral, desde donde llegas al resto "
              "de funciones:")
    steps(doc, [
        "① “Educación”: contenido para tu embarazo.",
        "② “Signos de alarma”: reportar síntomas.",
        "③ “Visitas domiciliarias”: historial de visitas a tu domicilio.",
        "④ “Mi perfil”: tus datos, tu FUM y tus preferencias.",
    ])
    shot(doc, "MENU", "Menú lateral de navegación.")
    section_gap(doc)

    # 5. Citas
    h1(doc, "5. Mis citas")
    para(doc, "Aquí llevas el control de tus citas de control prenatal (la meta son 8 controles).")
    steps(doc, [
        "① Pestañas “Próximas” e “Historial” para alternar entre tus citas futuras y las ya realizadas.",
        "② “Controles prenatales”: muestra cuántos controles llevas de la meta de 8.",
    ])
    shot(doc, "A2", "Listado de citas de control prenatal.")

    h2(doc, "5.1 Ver el detalle de una cita")
    para(doc, "Toca cualquier cita de la lista para abrir su detalle. En esa ventana puedes ver la fecha, la "
              "hora, el profesional, el número de control y las indicaciones.")
    shot(doc, "A2b", "Ventana de detalle de la cita.") if os.path.exists(os.path.join(IMG, "A2b.png")) else None
    para(doc, "Desde el detalle, según el estado de la cita, puedes:")
    steps(doc, [
        "Tocar “Confirmar” para confirmar tu asistencia.",
        "Tocar “Reprogramar” o “Solicitar reprogramación” para pedir un cambio de día. Se abrirá una ventana "
        "donde eliges la nueva fecha, un horario disponible y escribes el motivo (mínimo 5 caracteres). "
        "Tu obstetra debe aprobar la solicitud; te avisaremos cuando lo haga.",
    ])
    note(doc, "Una solicitud de reprogramación queda pendiente hasta que tu obstetra la apruebe. Mientras "
              "tanto, tu cita original sigue vigente.")
    section_gap(doc)

    # 6. Tratamiento
    h1(doc, "6. Mi tratamiento")
    para(doc, "En esta pantalla registras la toma de tus vitaminas y pastillas, y ves tu constancia.")
    steps(doc, ["① “Mi adherencia”: muestra tu porcentaje de cumplimiento y las dosis tomadas de los últimos días."])
    shot(doc, "A3", "Resumen de adherencia al tratamiento.")
    h2(doc, "6.1 Marcar un medicamento como tomado")
    steps(doc, ["① Cuando tomes una vitamina o pastilla, toca “Marcar como tomado”. El botón cambiará a "
                "“Tomado hoy” y tu porcentaje de cumplimiento se actualizará."])
    shot(doc, "A3b", "Botón “Marcar como tomado” en cada medicamento.")
    note(doc, "Tu obstetra puede ver tu constancia. Tomar tus suplementos a diario ayuda a prevenir la anemia "
              "y cuida a tu bebé.")
    section_gap(doc)

    # 7. Chat
    h1(doc, "7. Chat con mi obstetra")
    para(doc, "Puedes escribirle a tu obstetra para hacer preguntas o contarle cómo te sientes.")
    steps(doc, [
        "① Toca “Chat” en la barra inferior para abrir la conversación con tu obstetra.",
        "② Escribe tu mensaje en el campo inferior. ③ Toca el botón de enviar. Con el clip puedes adjuntar una foto.",
    ])
    shot(doc, "A4", "Conversación con la obstetra.")
    section_gap(doc)

    # 8. Educación
    h1(doc, "8. Educación")
    para(doc, "Encuentra artículos y recursos sencillos, elegidos según tu mes de embarazo.")
    steps(doc, [
        "① Buscador: escribe un tema para encontrar contenido.",
        "② “Mis semanas”: abre la calculadora de semanas de embarazo.",
        "③ “Signos de alarma”: acceso directo para reportar síntomas.",
    ])
    shot(doc, "A5", "Biblioteca de contenido educativo.")
    h2(doc, "8.1 Calcular mis semanas de embarazo")
    para(doc, "Al tocar “Mis semanas” se abre una calculadora. Elige la fecha de tu última regla (FUM) y toca "
              "“Ver mis semanas de embarazo”; verás tu semana actual, el trimestre y la fecha probable de parto.")
    shot(doc, "A5b", "Calculadora “¿En qué semana estoy?”.")
    section_gap(doc)

    # 9. Signos de alarma
    h1(doc, "9. Signos de alarma y emergencia")
    para(doc, "Si algo te preocupa, puedes reportar uno o varios síntomas a tu obstetra.")
    steps(doc, [
        "Marca los síntomas que presentas (están agrupados en “Durante el Embarazo”, “Durante el Parto” y "
        "“Después del Parto”).",
        "Si quieres, escribe más detalles en “Información adicional”.",
        "① Toca “Enviar alerta a mi obstetra”. Recibirás una confirmación y tu obstetra será notificada.",
    ])
    shot(doc, "A6", "Pantalla para reportar signos de alarma.")
    note(doc, "Si tienes sangrado, dolor de cabeza muy fuerte, pérdida de líquido, fiebre o tu bebé se mueve "
              "menos, usa el botón de Emergencia del inicio y acude de inmediato al establecimiento de salud. "
              "Teléfono: 083 – 421800.", "alerta")
    section_gap(doc)

    # 10. Perfil
    h1(doc, "10. Mi perfil")
    para(doc, "Desde tu perfil editas tus datos, eliges cómo recibir avisos y puedes repetir el recorrido guiado.")
    steps(doc, [
        "① “Mis datos y fecha de última regla”: edita tu información y registra tu FUM.",
        "② “Notificaciones”: elige por qué canales recibir avisos.",
        "③ “Conoce tu app”: repite el recorrido guiado cuando quieras.",
    ])
    shot(doc, "A7", "Pantalla de perfil de la gestante.")

    h2(doc, "10.1 Editar mis datos y mi FUM")
    para(doc, "Al tocar “Mis datos y fecha de última regla” se abre un formulario con tus datos. Los campos "
              "con asterisco (*) son obligatorios:")
    steps(doc, [
        "Completa “Nombres *”, “Apellidos *”, “Teléfono” y “Correo Electrónico”.",
        "Elige tu “Fecha de Nacimiento *”.",
        "Elige tu “Fecha de tu última regla (FUM) *”: es el primer día de tu última menstruación. Con ella se "
        "calculan tus semanas y tu cronograma de 8 controles.",
        "① Toca “Guardar Datos” para confirmar (o “Cancelar” para salir sin cambios).",
    ])
    shot(doc, "A7b", "Formulario “Modificar Perfil y FUM”.")

    h2(doc, "10.2 Preferencias de notificación")
    para(doc, "Al tocar “Notificaciones” eliges por qué canales quieres recibir recordatorios y alertas:")
    steps(doc, [
        "Activa o desactiva “Notificaciones en la app”, “SMS” y “WhatsApp” según prefieras.",
        "① Toca “Guardar” para confirmar tus preferencias.",
    ])
    shot(doc, "A7c", "Ventana “Preferencias de notificación”.")
    note(doc, "Las alertas clínicas urgentes siempre se enviarán por seguridad, aunque desactives algún canal.")
    section_gap(doc)

    # 11. Troubleshooting
    h1(doc, "11. Solución de problemas")
    for q, a in [
        ("No puedo iniciar sesión", "Verifica que tu DNI y contraseña sean correctos. Si los olvidaste, usa “¿Olvidaste tu contraseña?” o pide ayuda a tu obstetra."),
        ("No veo mi semana de embarazo", "Asegúrate de haber registrado tu fecha de última regla (FUM) en “Mi perfil” → “Mis datos y fecha de última regla”."),
        ("La app no carga o se ve en blanco", "Revisa tu conexión a internet, cierra la aplicación y vuelve a abrirla."),
        ("No recibo los recordatorios", "Entra a “Mi perfil” → “Notificaciones” y activa los canales que prefieras."),
    ]:
        h3(doc, q); para(doc, a)

    # 12. FAQ
    h1(doc, "12. Preguntas frecuentes")
    for q, a in [
        ("¿Cuántos controles prenatales debo tener?", "La meta son 8 controles durante el embarazo. La app te ayuda a llevar la cuenta."),
        ("¿Para qué sirve marcar mis pastillas?", "Para llevar tu constancia y que tu obstetra sepa que las estás tomando bien."),
        ("¿Puedo usar la app sin internet?", "Algunas acciones se guardan y se sincronizan cuando recuperas la conexión, pero se recomienda usarla con internet."),
        ("¿Mis datos están seguros?", "Sí. Tu información de salud solo es visible para el personal de salud autorizado."),
    ]:
        h3(doc, q); para(doc, a)

    # 13. Glosario
    h1(doc, "13. Glosario")
    for t, d in [
        ("FUM", "Fecha de la Última Menstruación. Con ella se calculan las semanas de tu embarazo."),
        ("FPP", "Fecha Probable de Parto."),
        ("Control prenatal", "Cita médica para revisar tu salud y la de tu bebé durante el embarazo."),
        ("Adherencia", "Qué tan constante eres tomando tus vitaminas y pastillas."),
        ("Signo de alarma", "Síntoma que puede indicar un problema y por el que debes buscar atención."),
        ("Modal / ventana", "Una ventana que se abre sobre la pantalla para una tarea concreta (por ejemplo, editar tus datos)."),
    ]:
        p = doc.add_paragraph(); p.add_run(f"{t}: ").bold = True; p.add_run(d)

    # 14. Contacto
    h1(doc, "14. Soporte y contacto")
    para(doc, "Si necesitas ayuda con la aplicación, comunícate con tu obstetra desde el chat de la app. Para "
              "consultas presenciales o emergencias, acude a tu establecimiento de salud.")
    p = doc.add_paragraph(); p.add_run("Centro de Salud Talavera\n").bold = True
    p.add_run("Andahuaylas, Apurímac\nTeléfono: 083 – 421800")

    cp = doc.core_properties
    cp.author = AUTHOR
    cp.title = "Manual de Usuario VITMATERNA — Guía para la Gestante"
    cp.category = "Manual de Usuario"
    doc.save(OUT)
    print(f"✅ DOCX: {OUT}  ·  figuras: {FIG['n']}  ·  autor: {AUTHOR}")


if __name__ == "__main__":
    build()
